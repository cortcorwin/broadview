import requests
import markdown
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from bs4 import BeautifulSoup
from bs4 import NavigableString
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

# Constants for formatting
DEFAULT_LINE_SPACING = 1.5
LIST_INDENT_PER_LEVEL = 44

def ultimate_word_doc_generator_v3(markdown_input: str, file_name: str):
    """
    Generates a Word doc from scratch with a custom header and footer.

    Args:
        markdown_input (string): Markdown that will be turned into the word document
        file_name (string): Name of the document
    """
    # Create a blank document
    doc = Document()
    print("Blank document created.")

    # Set default margins (optional, adjust if needed)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Add Header to All Pages ---
    header_url = "https://i.ibb.co/6cnL8Xxd/bv-header.png"
    header_response = requests.get(header_url)
    if header_response.status_code == 200:
        header_stream = io.BytesIO(header_response.content)
        header = doc.sections[0].header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = header_para.add_run()
        run.add_picture(header_stream, width=Inches(3.5))
        print("Header image added to first section.")
    else:
        print(f"Failed to download header image: {header_response.status_code}")

    # --- Add Space at the Top of the First Page ---
    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(72)  # 1 inch = 72 points

    # --- Extract Cover Page Data and Proposal Content in One Pass ---
    cover_data = {}
    proposal_content = []
    in_cover_section = False

    markdown_lines = markdown_input.split('\n')
    for line in markdown_lines:
        if line.strip() == '# Cover Page':
            in_cover_section = True
        elif in_cover_section and line.startswith('# '):
            in_cover_section = False
            proposal_content.append(line)
        elif in_cover_section and line.strip().startswith('**'):
            match = re.match(r'\*\*(.*?):\*\*\s*(.*)', line.strip())
            if match:
                key, value = match.groups()
                cover_data[key.strip()] = value.strip()
        elif not in_cover_section:
            proposal_content.append(line)
    print("Cover page data extracted:", cover_data)

    # --- Add Cover Page Content ---
    # Add a blank line with 28pt font size to push content down (existing spacer)
    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(28)

    # Add cover page data
    for key, value in cover_data.items():
        p = doc.add_paragraph(f"{key}: {value}")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(12)

    # Add placeholder for customer logo
    p = doc.add_paragraph("Customer Logo HERE")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(48)

    # --- Add Footer Image to Cover Page Footer ---
    footer_url = "https://i.ibb.co/Gf1stFFT/BV-footer.png"
    image_response = requests.get(footer_url)
    if image_response.status_code == 200:
        image_stream = io.BytesIO(image_response.content)
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs is not None and len(footer.paragraphs) > 0 else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer_para.add_run()
        run.add_picture(image_stream, width=Inches(6.5))
        print("Footer image added to cover page footer.")
    else:
        print(f"Failed to download footer image: {image_response.status_code}")

    # Add section break to separate cover page from proposal
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    for section in doc.sections[1:]:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].clear()

    # --- Add Page Numbers to Footer on All Pages ---
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer_para.add_run()
        field_code = OxmlElement('w:fldChar')
        field_code.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.text = "PAGE"
        field_code_end = OxmlElement('w:fldChar')
        field_code_end.set(qn('w:fldCharType'), 'end')
        run._element.append(field_code)
        run._element.append(instr_text)
        run._element.append(field_code_end)
        for r in footer_para.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(10)

    # --- Process Proposal Content ---
    proposal_markdown = '\n'.join(proposal_content)
    proposal_html = markdown.markdown(proposal_markdown, extensions=['extra', 'tables'])
    proposal_soup = BeautifulSoup(proposal_html, "html.parser")
    print("Proposal content converted to HTML.")

    # Set styles for proposal content
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0, 0, 0)

    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Arial'
    font_h1.size = Pt(18)
    font_h1.color.rgb = RGBColor(0, 0, 0)

    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Arial'
    font_h2.size = Pt(16)

    style_h3 = doc.styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = 'Arial'
    font_h3.size = Pt(14)

    # Helper functions for proposal content
    def parse_inline(soup_element, paragraph):
        if soup_element.contents is not None and len(soup_element.contents) > 0:
            for content in soup_element.contents:
                if isinstance(content, NavigableString):
                    text = content.strip()
                    if text:
                        run = paragraph.add_run(text)
                        run.font.name = 'Arial'
                        run.font.color.rgb = RGBColor(0, 0, 0)
                elif content.name in ["strong", "b"]:
                    text = content.get_text(strip=True)
                    if text:
                        run = paragraph.add_run(text)
                        run.font.name = 'Arial'
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    parse_inline(content, paragraph)

    def parse_table(table_element, doc):
        rows = table_element.find_all('tr')
        if not rows:
            return
        num_rows = len(rows)
        num_cols = len(rows[0].find_all(['th', 'td'])) if rows else 0
        if num_rows == 0 or num_cols == 0:
            return

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                cell_text = cell.get_text(strip=True)
                table.cell(i, j).text = cell_text
                if i == 0 or cell.find(['strong', 'b']):
                    for run in table.cell(i, j).paragraphs[0].runs:
                        run.font.name = 'Arial'
                        run.bold = True
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        print("Table added to document with space after.")

    def parse_element(element, doc):
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            heading = doc.add_heading(level=level)
            parse_inline(element, heading)
            if level in [1, 2, 3]:
                heading.paragraph_format.space_after = Pt(12)
        elif element.name == "p":
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para.paragraph_format.line_spacing = DEFAULT_LINE_SPACING
            parse_inline(element, para)
        elif element.name in ["ul", "ol"]:
            parse_list(element, level=0)
        elif element.name == "table":
            parse_table(element, doc)

    def parse_list(list_element, level):
        list_style = "List Number" if list_element.name == "ol" else "List Bullet"
        items = list_element.find_all("li", recursive=False)
        if items is not None and len(items) > 0:
            for li in items:
                nested_list = li.find(["ul", "ol"], recursive=False)
                if nested_list:
                    nested_list.extract()
                paragraph = doc.add_paragraph(style=list_style)
                paragraph.paragraph_format.left_indent = Pt(LIST_INDENT_PER_LEVEL * (level + 1))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.paragraph_format.line_spacing = DEFAULT_LINE_SPACING
                parse_inline(li, paragraph)
                if nested_list:
                    parse_list(nested_list, level + 1)

    # Process proposal content
    for element in proposal_soup.find_all(recursive=False):
        parse_element(element, doc)
    print("Proposal content processed.")

    # --- Add Page Break and Notes Image ---
    doc.add_page_break()
    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(50)
    notes_url = "https://i.ibb.co/vxhq0P8s/notes-template.png"
    image_response = requests.get(notes_url)
    if image_response.status_code == 200:
        image_stream = io.BytesIO(image_response.content)
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(image_stream, width=Inches(6.5))
        print("Notes image added successfully.")
    else:
        print(f"Failed to download notes image: {image_response.status_code}")
        p = doc.add_paragraph("Notes image could not be loaded due to a download error.")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 0, 0)

    # --- Save the Document ---
    doc.save(file_name)
    print(f"Document saved as {file_name}")
    return {"filename": file_name}
